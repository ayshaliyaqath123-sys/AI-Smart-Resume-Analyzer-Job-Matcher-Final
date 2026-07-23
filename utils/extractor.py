"""
extractor.py
Handles pulling raw text out of uploaded resume files (PDF or DOCX).
"""
import os
import pdfplumber
import docx


def extract_text(filepath: str) -> str:
    """Extract plain text from a .pdf or .docx file. Returns '' on failure."""
    ext = os.path.splitext(filepath)[1].lower()

    if ext == ".pdf":
        return _extract_from_pdf(filepath)
    elif ext == ".docx":
        return _extract_from_docx(filepath)
    else:
        raise ValueError(f"Unsupported file type: {ext}. Please upload a PDF or DOCX file.")


def _extract_from_pdf(filepath: str) -> str:
    text_parts = []
    with pdfplumber.open(filepath) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
    return "\n".join(text_parts)


def _extract_from_docx(filepath: str) -> str:
    document = docx.Document(filepath)
    paragraphs = [p.text for p in document.paragraphs if p.text.strip()]

    # Also grab text inside tables (some resumes use table layouts)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    paragraphs.append(cell.text)

    return "\n".join(paragraphs)
